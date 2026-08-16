# -*- coding: utf-8 -*-
"""
阅读理解题目追加工具
====================
在改编版 docx 尾部追加组卷网配套题目（含答案与解析），自动应用排版规则：
  - 标题「二、题型名」+ 题干指令
  - 正文首行缩进 2 字符、中文微软雅黑 / 英文 Arial、1.5 倍行距
  - 空处 ____NN____ 自动转成「不间断空格+数字」下划线格式，整条线对齐
  - 选词填空方框单词用带边框表格
  - 阅读单选：题干加粗、每个选项单独一行
  - 阅读问答：每题后跟 49 个下划线的答案横线、【答案】每题单独一行
  - 答案块【答案】【导语】【详解】整体 F2F2F2 底纹

用法：
  python 阅读理解题目追加工具.py                                  # 交互式向导
  python 阅读理解题目追加工具.py <docx路径> <spec.json>            # 按 spec 追加
  python 阅读理解题目追加工具.py --example                        # 生成 spec 示例
  python 阅读理解题目追加工具.py --help

spec.json 字段说明（均可省略，省略则跳过对应部分）：
{
  "type":        "选词填空 | 7选5 | 语法填空 | 首字母填空 | 阅读问答 | 阅读单选",
  "instruction": "题干指令（照抄原题）",
  "wordbank":    "选词填空的方框单词，空格分隔，如 'change quiet nervous'",
  "passage":     ["短文第1段", "第2段", "..."]，空处用 ____NN____ 标记，
  "options":     ["A．...", "B．...", ...]，7选5 的 A-G 选项，
  "questions":   [{"stem": "16．题干", "options": ["A．...", "B．..."]}, ...]，
                 阅读单选/阅读问答的题目（阅读问答 options 省略）
  "answers":     "答案一行，如 '51．stars　52．when'"，阅读问答请用 "answer_lines"
  "answer_lines":["157．About 2,000 years ago.", "158．..."]，阅读问答答案每题一行
  "summary":     "导语文字",
  "details":     ["【详解】第1条", "第2条", "..."]
}
"""

import sys, re, os, json
import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

SHADE = "F2F2F2"   # 答案块底纹
BAR = 49           # 阅读问答答案横线长度（下划线个数）
INDENT = "425"     # 首行缩进（twips，约 2 字符）

TYPE_TITLE = {
    "选词填空": "选词填空",
    "7选5": "七选五",
    "语法填空": "语法填空",
    "首字母填空": "首字母填空",
    "阅读问答": "阅读问答",
    "阅读单选": "阅读理解",
}


def set_fonts(run):
    run.font.name = "Arial"
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), "微软雅黑")


def _style(p, shade=None, indent=False):
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pPr = p._element.get_or_add_pPr()
    if shade:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), shade)
        pPr.append(shd)
    if indent:
        ind = OxmlElement("w:ind")
        ind.set(qn("w:firstLine"), INDENT)
        pPr.append(ind)


def add_para(doc, text, shade=None, bold=False, indent=False):
    """新增一个段落；空处 ____NN____ 自动转为下划线格式（数字带下划线）。"""
    p = doc.add_paragraph()
    _style(p, shade, indent)
    if re.search(r"____\d+____", text):
        pos = 0
        for m in re.finditer(r"(____)(\d+)(____)", text):
            if text[pos:m.start()]:
                r = p.add_run(text[pos:m.start()]); set_fonts(r); r.bold = bold
            seg = " " * 4 + m.group(2) + " " * 4
            r = p.add_run(seg); set_fonts(r); r.bold = bold; r.font.underline = True
            pos = m.end()
        if text[pos:]:
            r = p.add_run(text[pos:]); set_fonts(r); r.bold = bold
    else:
        r = p.add_run(text); set_fonts(r); r.bold = bold
    return p


def add_wordbank(doc, words):
    """选词填空方框：单格带边框表格。"""
    tbl = doc.add_table(rows=1, cols=1)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement("w:" + edge)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tbl._tbl.tblPr.append(borders)
    cp = tbl.rows[0].cells[0].paragraphs[0]
    cp.paragraph_format.line_spacing = 1.5
    cp.paragraph_format.space_after = Pt(0)
    r = cp.add_run(words); set_fonts(r)


def add_bar(doc, n=BAR):
    """阅读问答答案横线：一整行下划线。"""
    p = doc.add_paragraph()
    _style(p)
    r = p.add_run("_" * n); set_fonts(r)


def append_from_spec(docx_path, spec):
    if not os.path.isfile(docx_path):
        raise FileNotFoundError("找不到文件: %s" % docx_path)
    doc = docx.Document(docx_path)
    qtype = spec.get("type", "")
    if qtype not in TYPE_TITLE:
        raise ValueError("type 必须是 %s 之一" % "/".join(TYPE_TITLE))

    add_para(doc, "")
    add_para(doc, "二、" + TYPE_TITLE[qtype], bold=True)
    if spec.get("instruction"):
        add_para(doc, spec["instruction"])
    add_para(doc, "")

    # 方框单词
    if spec.get("wordbank"):
        add_wordbank(doc, spec["wordbank"])
        add_para(doc, "")

    # 短文
    for para in spec.get("passage", []):
        add_para(doc, para, indent=True)
    if spec.get("passage"):
        add_para(doc, "")

    # 7选5 / 选词填空 A-E 的选项列表
    for opt in spec.get("options", []):
        add_para(doc, opt)

    # 阅读单选 / 阅读问答的题目
    questions = spec.get("questions", [])
    for q in questions:
        stem = q.get("stem", "")
        if not stem:
            continue
        add_para(doc, stem, bold=(qtype == "阅读单选"))
        for o in q.get("options", []):
            add_para(doc, o)
        if qtype == "阅读问答":
            add_bar(doc)
    if questions:
        add_para(doc, "")

    # 答案块
    if spec.get("answer_lines"):
        for line in spec["answer_lines"]:
            add_para(doc, line, shade=SHADE)
    elif spec.get("answers"):
        add_para(doc, "【答案】" + spec["answers"], shade=SHADE)
    if spec.get("summary"):
        add_para(doc, "【导语】" + spec["summary"], shade=SHADE, indent=True)
    for i, d in enumerate(spec.get("details", [])):
        prefix = "【详解】" if i == 0 else ""
        add_para(doc, prefix + d, shade=SHADE, indent=True)

    doc.save(docx_path)
    return doc.paragraphs


def _read_lines(prompt):
    print(prompt)
    print("（多行请逐行粘贴，输入一个空行结束；不需要则直接回车）")
    lines = []
    while True:
        line = input()
        if line.strip() == "":
            break
        lines.append(line)
    return lines


def interactive():
    print("== 阅读理解题目追加工具 · 交互模式 ==")
    path = input("docx 文件路径：").strip().strip('"')
    print("题型：1=选词填空 2=7选5 3=语法填空 4=首字母填空 5=阅读问答 6=阅读单选")
    t = input("选择（1-6）：").strip()
    qtype = {1: "选词填空", 2: "7选5", 3: "语法填空",
             4: "首字母填空", 5: "阅读问答", 6: "阅读单选"}.get(int(t), "选词填空")

    spec = {"type": qtype}
    inst = input("题干指令（无则回车）：").strip()
    if inst:
        spec["instruction"] = inst
    if qtype == "选词填空":
        wb = input("方框单词（空格分隔，无则回车）：").strip()
        if wb:
            spec["wordbank"] = wb
    if qtype in ("阅读单选", "阅读问答"):
        qs = []
        print("逐题输入：先粘贴题干行，再逐行粘贴选项，输入空行表示该题结束，再输入空行表示全部结束")
        while True:
            stem = input("题干（空行结束）：").strip()
            if not stem:
                break
            opts = []
            while True:
                o = input("  选项（空行结束）：").strip()
                if not o:
                    break
                opts.append(o)
            qs.append({"stem": stem, "options": opts})
        spec["questions"] = qs
    else:
        passage = _read_lines("请粘贴短文（空处用 ____NN____ 标记）")
        spec["passage"] = passage
        if qtype == "7选5":
            opts = _read_lines("请粘贴 A-G 选项（每行一个）")
            spec["options"] = opts
    if qtype == "阅读问答":
        spec["answer_lines"] = _read_lines("请逐行粘贴答案（每题一行）")
    else:
        ans = input("答案一行（如 51．stars　52．when，无则回车）：").strip()
        if ans:
            spec["answers"] = ans
    sm = input("导语（无则回车）：").strip()
    if sm:
        spec["summary"] = sm
    details = _read_lines("请粘贴详解（每题一条；第一条可写【详解】或直接写 16．…，自动加【详解】前缀）")
    spec["details"] = details

    append_from_spec(path, spec)
    print("完成：已追加到", path)


def write_example():
    ex = {
        "type": "阅读单选",
        "instruction": "",
        "passage": [
            "Our five senses are amazing gifts that help us understand and enjoy the world around us. They allow us to see, hear, smell, taste, and touch.",
            "Our senses also help us deal with stress. When we feel anxious, taking a deep breath and smelling something pleasant can help us relax.",
        ],
        "questions": [
            {"stem": "16．What is the main idea of the passage?",
             "options": ["A．The five senses are not very important.",
                         "B．The five senses help us enjoy life and deal with stress.",
                         "C．Only vision and hearing are useful.",
                         "D．We don’t need our senses in daily life."]},
            {"stem": "17．Which sense can warn us of danger like smoke or gas?",
             "options": ["A．Vision.", "B．Hearing.", "C．Taste.", "D．Smell."]},
        ],
        "answers": "16．B　17．D",
        "summary": "本文是一篇说明文，介绍人的五种感官各自的功能，说明感官丰富我们的生活，还可以帮助我们缓解压力。",
        "details": ["16．通读全文可知，选项B概括最全面准确。",
                    "17．根据第二段中关于嗅觉的描述，可知是嗅觉。"],
    }
    with open("spec示例.json", "w", encoding="utf-8") as f:
        json.dump(ex, f, ensure_ascii=False, indent=2)
    print("已生成 spec示例.json，请参考后填写自己的题目内容。")


def main():
    args = sys.argv[1:]
    if not args or args[0] in ("--help", "-h"):
        print(__doc__)
        return
    if args[0] == "--example":
        write_example()
        return
    if args[0] == "--interactive":
        interactive()
        return
    if len(args) >= 2:
        docx_path, spec_path = args[0], args[1]
        with open(spec_path, encoding="utf-8") as f:
            spec = json.load(f)
        append_from_spec(docx_path, spec)
        print("完成：已追加到", docx_path)
        return
    interactive()


if __name__ == "__main__":
    main()
