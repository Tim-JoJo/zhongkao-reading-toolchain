"""Word 文档导出。

生成格式：标题 → 说明 → 正文 → 选择题 → Answer Key
样式：Arial 12pt 英文，微软雅黑中文，行间距 1.5 倍，标题 20pt。
"""

from __future__ import annotations

import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    try:
        # python-docx >= 1.2.0：WD_COLOR_INDEX 位于 docx.enum.text
        from docx.enum.text import WD_COLOR_INDEX
    except ImportError:
        # python-docx < 1.2.0：WD_COLOR_INDEX 位于 docx.enum.dml
        from docx.enum.dml import WD_COLOR_INDEX
except ImportError:
    Document = None


def _set_run_font(run, size: int = 12, bold: bool = False, italic: bool = False, east_asia: str = "微软雅黑"):
    """统一设置 run 的字体和大小。east_asia 为中文/全角字符字体。"""
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


# 排序题事件标号：①②③④（用苹方-简字体渲染）
ORDERING_MARK_RE = re.compile(r"^([①②③④])([\.\s].*)$")


def _set_mark_font(run, size: int = 12, bold: bool = False):
    """排序题标号 ①②③④：ascii/hAnsi/eastAsia/cs 全部设为苹方-简，
    保证 Word 对全角符号按苹方渲染（仅设 eastAsia 时 ① 仍走 Arial）。"""
    run.font.size = Pt(size)
    run.bold = bold
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn("w:" + attr), "苹方-简")


def _get_question_type_label(q: dict) -> str:
    """从题目 dict 推导题型标签字符串。

    优先用 type 字段；vocabulary_or_detail 二选一类型靠 code（蓝图编号
    V-xx / D-xx）进一步区分 VOCABULARY 与 DETAIL。
    """
    t = q.get("type", "")
    code = q.get("code", "")
    if t == "writing_technique":
        return "WRITING TECHNIQUE"
    if t == "detail":
        return "DETAIL"
    if t == "vocabulary_or_detail":
        if code and code.startswith("V"):
            return "VOCABULARY"
        if code and code.startswith("D"):
            return "DETAIL"
        return "VOCABULARY / DETAIL"
    if t == "inference":
        return "INFERENCE"
    if t == "ordering":
        return "ORDERING"
    if t == "main_idea":
        return "MAIN IDEA"
    return ""


def run_export_docx(
    *,
    title: str,
    body: str,
    questions: list[dict],
    answer_key: list[str],
    output_path: str,
    explanations: list[str] | None = None,
) -> str:
    """将正文和选择题导出为 .docx 文档。

    explanations: 每题答案解析（与题目一一对应），渲染在 Answer Key 之后。
    """
    if Document is None:
        return "错误：需要安装 python-docx 库 (pip install python-docx)"

    try:
        doc = Document()

        # ── 默认样式：Arial 12pt + 微软雅黑 + 1.5 倍行距 ──
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Arial"
        font.size = Pt(12)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        pf = style.paragraph_format
        pf.line_spacing = 1.5

        # ── 标题（20pt） ──
        h = doc.add_heading("Reading Comprehension", level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in h.runs:
            run.font.name = "Arial"
            run.font.size = Pt(20)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        # ── 说明 ──
        inst = doc.add_paragraph()
        run = inst.add_run("Read the passage and choose the best answer for each question.")
        _set_run_font(run, size=10, italic=True)

        doc.add_paragraph()

        # ── 正文标题（20pt 加粗，居中） ──
        pt_para = doc.add_paragraph()
        pt_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pt_para.add_run(title)
        _set_run_font(run, size=20, bold=True)

        # ── 正文（12pt，1.5 倍行距） ──
        for para_text in body.split("\n\n"):
            p = doc.add_paragraph(para_text.strip())
            p.paragraph_format.first_line_indent = Cm(0.75)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.5
            for r in p.runs:
                r.font.name = "Arial"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        doc.add_paragraph()

        # ── 选择题 ──
        for i, q in enumerate(questions):
            # 题型标签（右对齐，小号加粗灰色，标在题干上方）
            type_label = _get_question_type_label(q)
            if type_label:
                tag_para = doc.add_paragraph()
                tag_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                tag_para.paragraph_format.space_before = Pt(16)
                tag_para.paragraph_format.space_after = Pt(0)
                tag_para.paragraph_format.line_spacing = 1.0
                run = tag_para.add_run(type_label)
                _set_run_font(run, size=8, bold=True)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.line_spacing = 1.5
            # 悬挂缩进：续行（排序题事件列表）缩进约 0.5cm，与首行 "4. Put" 的 P 大致对齐
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            stem = q.get("stem", "")
            is_ordering = q.get("type") == "ordering"
            # 题干首行加粗；续行（如排序题的事件列表）不加粗，续行之间用换行连接。
            # 排序题事件标号（①~④）用苹方-简字体，且续行换行用 <w:br/> 保证 Word 一行一个事件。
            parts = stem.split("\n")
            for j, line in enumerate(parts):
                if not line.strip():
                    continue
                if j == 0:
                    run = p.add_run(f"{i+1}. {line}")
                    _set_run_font(run, size=12, bold=True)
                else:
                    # 换行（w:br）
                    br_run = p.add_run()
                    _set_run_font(br_run, size=12, bold=False)
                    br = br_run._element.makeelement(qn("w:br"), {})
                    br_run._element.append(br)
                    # 事件行文本
                    m = ORDERING_MARK_RE.match(line.strip()) if is_ordering else None
                    if m:
                        mark_run = p.add_run(m.group(1))
                        _set_mark_font(mark_run, size=12, bold=False)
                        text_run = p.add_run(m.group(2))
                        _set_run_font(text_run, size=12, bold=False)
                    else:
                        run = p.add_run(line)
                        _set_run_font(run, size=12, bold=False)
            for opt in q.get("options", []):
                opt_text = opt.strip()
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.line_spacing = 1.5
                if is_ordering:
                    # 排序题选项里的 ①~④ 序列标号也用苹方-简，与题干事件标号一致
                    parts = re.split(r"([①②③④])", opt_text)
                    for seg in parts:
                        if not seg:
                            continue
                        if seg in "①②③④":
                            r = p.add_run(seg)
                            _set_mark_font(r, size=12, bold=False)
                        else:
                            r = p.add_run(seg)
                            _set_run_font(r, size=12, bold=False)
                else:
                    r = p.add_run(opt_text)
                    _set_run_font(r, size=12, bold=False)

        doc.add_paragraph()

        # ── Answer Key ──
        p = doc.add_paragraph()
        run = p.add_run("Answer Key:")
        _set_run_font(run, size=12, bold=True)
        p.paragraph_format.line_spacing = 1.5
        key_text = "   ".join(f"{i+1}. {a}" for i, a in enumerate(answer_key))
        key_para = doc.add_paragraph(key_text)
        key_para.paragraph_format.line_spacing = 1.5
        for r in key_para.runs:
            r.font.name = "Arial"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        # ── 答案解析（每题一条） ──
        if explanations:
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run("答案解析:")
            _set_run_font(run, size=12, bold=True)
            p.paragraph_format.line_spacing = 1.5
            for i, expl in enumerate(explanations):
                p = doc.add_paragraph()
                # 解析文本可能已自带编号（如 "1. 解析：…"），剥掉避免与自动编号重复成 "1. 1."
                expl_text = re.sub(r"^\s*\d+\s*[.．、]\s*", "", expl.strip())
                run = p.add_run(f"{i+1}. {expl_text}")
                _set_run_font(run, size=10, italic=True)
                p.paragraph_format.line_spacing = 1.5

        # ── 保存 ──
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        return f"文档已保存至：{output_path}"

    except Exception as e:
        return f"导出失败：{e}"


def run_export_report_docx(
    *,
    title: str,
    content: list[dict[str, str]],
    output_path: str | None = None,
) -> str:
    """导出报告 Word 文档（文章正文带段内标注）。

    Args:
        title: 报告标题
        content: 报告章节列表，每项 {"heading": "...", "paragraphs": ["...", ...]}，
                 paragraph 内可用 [HIGHLIGHT:color]...[/HIGHLIGHT] 标注
                 （color 取值 yellow / turquoise / pink）
        output_path: 输出文件路径（含 .docx 扩展名），不传则保存至 Downloads\文章报告

    Returns:
        str: 成功时返回保存路径，失败时返回错误信息
    """
    if Document is None:
        return "错误：需要安装 python-docx 库 (pip install python-docx)"

    try:
        doc = Document()

        # ── 默认样式 ──
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Arial"
        font.size = Pt(12)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        pf = style.paragraph_format
        pf.line_spacing = 1.5

        # ── 报告标题（20pt 加粗，居中） ──
        pt_para = doc.add_paragraph()
        pt_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pt_para.add_run(title)
        _set_run_font(run, size=20, bold=True)

        HIGHLIGHT_COLORS = {
            "yellow": WD_COLOR_INDEX.YELLOW,
            "turquoise": WD_COLOR_INDEX.TURQUOISE,
            "pink": WD_COLOR_INDEX.PINK,
        }

        for section in content:
            heading = (section.get("heading") or "").strip()
            if heading:
                h = doc.add_paragraph()
                run = h.add_run(heading)
                _set_run_font(run, size=14, bold=True)
                h.paragraph_format.space_before = Pt(12)

            for para_text in section.get("paragraphs", []):
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(0.75)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.5

                # 解析 [HIGHLIGHT:color]...[/HIGHLIGHT] 标注
                # 注意：正则不感知嵌套。嵌套标注（如 [HIGHLIGHT:pink][HIGHLIGHT:yellow]...）
                # 会残留外层标记文本，渲染结束后统一剥离，保证输出不含 [HIGHLIGHT 字样。
                rest = para_text
                while rest:
                    m = re.match(r"^(.*?)\[HIGHLIGHT:(\w+)\](.*?)\[/HIGHLIGHT\](.*)$", rest, re.S)
                    if not m:
                        if rest.strip():
                            run = p.add_run(rest)
                            _set_run_font(run)
                        break
                    plain, color, marked, rest = m.groups()
                    if plain.strip():
                        run = p.add_run(plain)
                        _set_run_font(run)
                    run = p.add_run(marked)
                    _set_run_font(run)
                    run.font.highlight_color = HIGHLIGHT_COLORS.get(color, WD_COLOR_INDEX.YELLOW)

                # 兜底：剥离任何残留的 HIGHLIGHT 标记文本（含嵌套场景）
                for run in p.runs:
                    if "[HIGHLIGHT:" in run.text or "[/HIGHLIGHT]" in run.text:
                        cleaned = re.sub(r"\[/?HIGHLIGHT:\w*\]", "", run.text)
                        if cleaned:
                            run.text = cleaned
                        else:
                            run._element.getparent().remove(run._element)

        # ── 保存 ──
        out = Path(output_path or str(Path(DEFAULT_REPORT_DIR) / f"{title}.docx"))
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out))
        return f"报告已保存至：{out}"

    except Exception as e:
        return f"导出失败：{e}"


# ── 默认输出目录 ──
DEFAULT_OUTPUT_DIR = str(Path.home() / "Downloads")
DEFAULT_ARTICLE_DIR = str(Path.home() / "Downloads" / "生成文章")
DEFAULT_REPORT_DIR = str(Path.home() / "Downloads" / "文章报告")


def run_export_article_docx(
    *,
    title: str,
    body: str,
    output_path: str,
) -> str:
    """导出文章（不含题目），供 Part 1 文章写作完成后使用。

    生成格式：标题 → 正文。
    """
    if Document is None:
        return "错误：需要安装 python-docx 库 (pip install python-docx)"

    try:
        doc = Document()

        # ── 默认样式 ──
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Arial"
        font.size = Pt(12)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        pf = style.paragraph_format
        pf.line_spacing = 1.5

        # ── 正文标题（20pt 加粗，居中） ──
        pt_para = doc.add_paragraph()
        pt_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pt_para.add_run(title)
        _set_run_font(run, size=20, bold=True)

        # ── 正文 ──
        for para_text in body.split("\n\n"):
            p = doc.add_paragraph(para_text.strip())
            p.paragraph_format.first_line_indent = Cm(0.75)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.5
            for r in p.runs:
                r.font.name = "Arial"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        # ── 保存 ──
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        return f"文档已保存至：{output_path}"

    except Exception as e:
        return f"导出失败：{e}"
